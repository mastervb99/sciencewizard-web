"""
Velvet Research - Report Generation Service

Connects to science_wizard_project agents for manuscript generation.
"""
import os
import sys
import uuid
import asyncio
import traceback
from pathlib import Path
from datetime import datetime
from typing import Optional, Callable
import json

# Add science_wizard_project to path
SCIENCE_WIZARD_PATH = os.getenv(
    "SCIENCE_WIZARD_PATH",
    "/Users/vafabayat/Dropbox/Financial/0ScienceWizard/science_wizard_project"
)
sys.path.insert(0, SCIENCE_WIZARD_PATH)


class ReportGenerator:
    """
    Generates manuscripts using science_wizard_project agents.
    """

    def __init__(self, upload_path: str, output_dir: str):
        self.upload_path = Path(upload_path)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.progress = 0.0
        self.status = "pending"
        self.error = None
        self.report_path = None

    async def generate(
        self,
        progress_callback: Optional[Callable[[float, str], None]] = None
    ) -> str:
        """
        Run the full generation pipeline.

        Returns path to generated report.
        """
        try:
            self.status = "processing"

            # Step 1: Find and load data files
            await self._update_progress(0.1, "Loading data files...", progress_callback)
            data_files, doc_files = self._categorize_files()

            if not data_files and not doc_files:
                raise ValueError("No valid files found in upload")

            # Step 2: Import and run agents
            await self._update_progress(0.2, "Initializing AI agents...", progress_callback)

            # Import agents from science_wizard_project
            from agents.data_ingestor import DataIngestor
            from agents.research_planner import ResearchPlanner
            from agents.code_generator import CodeGenerator
            from agents.report_writer import ReportWriter

            # Step 3: Ingest data
            await self._update_progress(0.3, "Analyzing data structure...", progress_callback)
            ingestor = DataIngestor()

            data_profile = None
            if data_files:
                data_profile = await asyncio.to_thread(
                    ingestor.ingest,
                    str(data_files[0])  # Primary data file
                )

            # Step 4: Extract research context from documents
            await self._update_progress(0.4, "Extracting research context...", progress_callback)
            research_context = ""
            if doc_files:
                research_context = self._extract_text_from_docs(doc_files)

            # Step 5: Plan analysis
            await self._update_progress(0.5, "Planning statistical analysis...", progress_callback)
            planner = ResearchPlanner()
            analysis_plan = await asyncio.to_thread(
                planner.plan,
                data_profile,
                research_context
            )

            # Step 6: Generate analysis code
            await self._update_progress(0.6, "Generating analysis code...", progress_callback)
            code_gen = CodeGenerator()
            analysis_results = await asyncio.to_thread(
                code_gen.generate_and_execute,
                analysis_plan,
                data_profile
            )

            # Step 7: Write report
            await self._update_progress(0.8, "Writing manuscript...", progress_callback)
            writer = ReportWriter()
            report_path = self.output_dir / f"manuscript_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

            await asyncio.to_thread(
                writer.write,
                analysis_results,
                str(report_path),
                research_context
            )

            # Step 8: Complete
            await self._update_progress(1.0, "Complete!", progress_callback)
            self.status = "completed"
            self.report_path = str(report_path)

            return str(report_path)

        except Exception as e:
            self.status = "failed"
            self.error = str(e)
            traceback.print_exc()
            raise

    async def generate_simple(
        self,
        progress_callback: Optional[Callable[[float, str], None]] = None
    ) -> str:
        """
        Simplified generation for MVP - uses basic document synthesis.

        For when full pipeline isn't available or for simple cases.
        """
        try:
            self.status = "processing"
            await self._update_progress(0.1, "Loading files...", progress_callback)

            data_files, doc_files = self._categorize_files()

            # For MVP, generate a template-based report
            await self._update_progress(0.3, "Analyzing content...", progress_callback)

            # Extract any text content
            content = ""
            for doc in doc_files:
                content += self._extract_text_from_docs([doc]) + "\n\n"

            await self._update_progress(0.5, "Generating manuscript...", progress_callback)

            # Use Anthropic API directly for simple generation
            import anthropic

            client = anthropic.Anthropic()
            response = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=8000,
                messages=[{
                    "role": "user",
                    "content": f"""Based on the following research materials, generate a complete academic manuscript with these sections:
1. Abstract
2. Introduction
3. Methods
4. Results
5. Discussion
6. References

Research materials:
{content[:15000]}

Generate a well-structured academic manuscript."""
                }]
            )

            manuscript_text = response.content[0].text

            await self._update_progress(0.8, "Formatting document...", progress_callback)

            # Create Word document
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading("Research Manuscript", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add manuscript content
            for line in manuscript_text.split('\n'):
                if line.strip():
                    if line.startswith('#'):
                        level = line.count('#')
                        text = line.lstrip('#').strip()
                        doc.add_heading(text, level=min(level, 3))
                    else:
                        para = doc.add_paragraph(line)
                        para.paragraph_format.first_line_indent = Inches(0.5)

            # Save
            report_path = self.output_dir / f"manuscript_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(str(report_path))

            await self._update_progress(1.0, "Complete!", progress_callback)
            self.status = "completed"
            self.report_path = str(report_path)

            return str(report_path)

        except Exception as e:
            self.status = "failed"
            self.error = str(e)
            traceback.print_exc()
            raise

    def _categorize_files(self) -> tuple[list, list]:
        """Categorize uploaded files into data and documents."""
        data_files = []
        doc_files = []

        for f in self.upload_path.iterdir():
            if f.is_file():
                ext = f.suffix.lower()
                if ext in {'.csv', '.xlsx', '.xls'}:
                    data_files.append(f)
                elif ext in {'.docx', '.pdf', '.txt'}:
                    doc_files.append(f)

        return data_files, doc_files

    def _extract_text_from_docs(self, doc_files: list) -> str:
        """Extract text content from document files."""
        text_parts = []

        for doc_path in doc_files:
            ext = doc_path.suffix.lower()

            if ext == '.txt':
                text_parts.append(doc_path.read_text(errors='ignore'))

            elif ext == '.docx':
                try:
                    from docx import Document
                    doc = Document(str(doc_path))
                    text_parts.append('\n'.join([p.text for p in doc.paragraphs]))
                except Exception:
                    pass

            elif ext == '.pdf':
                try:
                    import fitz  # PyMuPDF
                    with fitz.open(str(doc_path)) as pdf:
                        for page in pdf:
                            text_parts.append(page.get_text())
                except ImportError:
                    # Fall back if PyMuPDF not available
                    pass

        return '\n\n'.join(text_parts)

    async def _update_progress(
        self,
        progress: float,
        message: str,
        callback: Optional[Callable] = None
    ):
        """Update progress and call callback if provided."""
        self.progress = progress
        if callback:
            await asyncio.to_thread(callback, progress, message)


# Background job runner
_running_jobs = {}


async def run_generation_job(
    job_id: str,
    upload_path: str,
    output_dir: str,
    db_update_callback
):
    """
    Run report generation as background task.
    """
    generator = ReportGenerator(upload_path, output_dir)
    _running_jobs[job_id] = generator

    try:
        # Use simplified generation for MVP
        report_path = await generator.generate_simple(
            progress_callback=lambda p, m: db_update_callback(job_id, "processing", p)
        )
        await db_update_callback(job_id, "completed", 1.0, report_path=report_path)

    except Exception as e:
        await db_update_callback(job_id, "failed", generator.progress, error=str(e))

    finally:
        if job_id in _running_jobs:
            del _running_jobs[job_id]


def get_job_progress(job_id: str) -> Optional[float]:
    """Get current progress of running job."""
    if job_id in _running_jobs:
        return _running_jobs[job_id].progress
    return None
